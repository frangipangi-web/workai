from docx import Document

doc = Document('./templates/Datatilsynet_skabelon-til-databehandleraftale-engelsk.docx')

print("=" * 80)
print("TEMPLATE STRUCTURE ANALYSIS")
print("=" * 80)
print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}\n")

print("First 60 paragraphs:")
print("-" * 80)
for i, p in enumerate(doc.paragraphs[:60]):
    text = p.text.strip()[:75] if p.text.strip() else "[empty]"
    level = p.style.name if p.style else "None"
    print(f"{i:3d} [{level:20s}] {text}")

print("\n" + "=" * 80)
print("TABLE CELLS (first 5 tables)")
print("=" * 80)
for t_idx, table in enumerate(doc.tables[:5]):
    print(f"\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
    for r_idx, row in enumerate(table.rows[:3]):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text[:50] if cell.text else "[empty]"
            print(f"  [{r_idx},{c_idx}]: {text}")
