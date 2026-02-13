#!/usr/bin/env python3
"""
GDPR Data Processor Agreement Generator
Generates GDPR-compliant data processor agreements with health data detection and country-specific adjustments.
"""

import json
import re
from datetime import datetime
from dataclasses import dataclass
from typing import Dict, List, Set, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@dataclass
class AgreementData:
    """Input data for DPA generation"""
    controller_name: str
    controller_cvr: str
    controller_address: str
    processor_name: str
    processor_cvr: str
    processor_address: str
    service_name: str
    data_types: List[str]
    storage_location: str
    country: str = "Denmark"  # Denmark, Norway, Germany
    execution_date: str = None
    additional_notes: str = ""
    data_categories: str = ""
    purposes: str = ""
    retention_period: str = "Duration of contract+"
    sub_processors: List[Dict] = None


class HealthDataDetector:
    """Detects health data and special category data in provided information"""
    
    HEALTH_KEYWORDS = {
        # Medical/Clinical
        'medical', 'doctor', 'physician', 'nurse', 'clinic', 'hospital', 'treatment',
        'diagnosis', 'diagnose', 'patient', 'prescription', 'medication', 'drug',
        'therapy', 'therapeutic', 'counseling', 'mental health', 'psychiatry', 'psychologist',
        'physiotherapy', 'physical therapy', 'rehabilitation', 'recovery',
        'surgery', 'surgical', 'operation', 'procedure',
        'symptom', 'condition', 'illness', 'disease', 'disorder', 'syndrome',
        'vaccine', 'vaccination', 'immunization',
        'genetic', 'dna', 'genome', 'biometric', 'fingerprint', 'facial recognition',
        'disability', 'disabled', 'impairment', 'handicap',
        # Health Services
        'healthcare', 'health center', 'health professional', 'practitioner',
        'pharmacy', 'dentist', 'optometrist', 'chiropractor',
        # Health Insurance
        'health insurance', 'medical insurance', 'disability insurance',
        'workers compensation', 'health plan', 'insurance claim',
        'health savings', 'hsa', 'fsa',
        # Health Records
        'medical history', 'patient records', 'clinical records', 'health records',
        'ehr', 'emr', 'health status',
        # Regulatory
        'hipaa', 'article 9', 'special category', 'protected health information', 'phi'
    }
    
    MAYBE_KEYWORDS = {
        'fitness', 'health', 'wellness', 'insurance', 'leave', 'absence',
        'accommodation', 'accessibility', 'sensitive'
    }
    
    POSITIVE_HEALTH_INDICATORS = [
        'health data', 'special category', 'sensitive data', 'medical',
        'patient', 'diagnosis', 'treatment', 'medication', 'health record'
    ]
    
    @staticmethod
    def detect_health_data(data_types: List[str], additional_text: str = "") -> Tuple[bool, str]:
        """
        Detect if provided data includes health data.
        Returns (is_health_data, reasoning_message)
        """
        combined_text = " ".join(data_types) + " " + additional_text
        combined_lower = combined_text.lower()
        
        # Check for definitive health indicators
        health_count = sum(1 for keyword in HealthDataDetector.HEALTH_KEYWORDS 
                          if keyword in combined_lower)
        
        if health_count >= 2:
            return True, "Multiple health-related indicators detected"
        
        if health_count == 1:
            # Check if it's a definitive indicator or possible false positive
            for keyword in HealthDataDetector.POSITIVE_HEALTH_INDICATORS:
                if keyword in combined_lower:
                    return True, f"Health data indicator found: '{keyword}'"
        
        maybe_count = sum(1 for keyword in HealthDataDetector.MAYBE_KEYWORDS 
                         if keyword in combined_lower)
        
        if maybe_count >= 3:
            return True, "Possible health data context detected (confirm with user)"
        
        return False, "No health data indicators detected"


class CountryModifier:
    """Applies country-specific modifications and text"""
    
    DEFAULT_CLAUSES = {
        "preamble": "For the purposes of Article 28(3) of Regulation (EU) 2016/679 (the GDPR)",
        "breach_notification": "without undue delay and, in any case, not later than 72 hours",
        "sub_processor_notice": "at least 30 days"
    }
    
    COUNTRY_MODIFICATIONS = {
        "Denmark": {
            "authority": "Datatilsynet (Danish Data Protection Authority)",
            "authority_url": "https://www.datatilsynet.dk",
            "references": ["GDPR Article 28(3)"],
            "preamble_addition": "",
            "breach_notification": "without undue delay and, in any case, not later than 72 hours",
            "sub_processor_notice": "at least 30 days"
        },
        "Norway": {
            "authority": "Datatilsynet (Norwegian Data Protection Authority)",
            "authority_url": "https://www.datatilsynet.no",
            "references": ["Section 2.7, Chapter 2 of Norwegian Personal Data Act", "GDPR Article 28(3)"],
            "preamble_addition": "and in accordance with Section 2.7, Chapter 2 of the Norwegian Personal Data Act",
            "breach_notification": "immediately and, in any case, not later than 48 hours",
            "sub_processor_notice": "at least 30 days",
            "data_residency_note": "Personal data shall be processed primarily within Norwegian territory or the EEA"
        },
        "Germany": {
            "authority": "BfDI (Bundesbeauftragter für Datenschutz und Informationsfreiheit)",
            "authority_url": "https://www.bfdi.bund.de",
            "references": ["Section 7 Federal Data Protection Act (BDSG)", "GDPR Article 28(3)"],
            "preamble_addition": "and Section 7 of the Federal Data Protection Act (BDSG)",
            "breach_notification": "immediately and, in any case, not later than 24 hours",
            "sub_processor_notice": "at least 45 days",
            "requires_audit": True,
            "audit_frequency": "Annual",
            "signature_requirement": "Original signatures of authorized representatives"
        }
    }
    
    @staticmethod
    def get_modification(country: str, key: str) -> str:
        """Get country-specific text modification"""
        if country not in CountryModifier.COUNTRY_MODIFICATIONS:
            country = "Denmark"
        
        modifications = CountryModifier.COUNTRY_MODIFICATIONS[country]
        return modifications.get(key, CountryModifier.DEFAULT_CLAUSES.get(key, ""))


class DPAGenerator:
    """Generates GDPR Data Processor Agreement document"""
    
    def __init__(self, data: AgreementData):
        self.data = data
        self.doc = Document()
        self.health_data_detected = False
        self.health_data_reason = ""
        
        # Detect health data
        self._detect_health_data()
        
        # Apply country defaults
        if not data.execution_date:
            data.execution_date = datetime.now().strftime("%B %d, %Y")
    
    def _detect_health_data(self):
        """Detect health data in the provided information"""
        detector = HealthDataDetector()
        is_health, reason = detector.detect_health_data(
            self.data.data_types,
            self.data.additional_notes + " " + self.data.data_categories
        )
        self.health_data_detected = is_health
        self.health_data_reason = reason
    
    def _add_heading(self, text: str, level: int = 1):
        """Add a heading to the document"""
        p = self.doc.add_paragraph(text, style=f'Heading {level}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        return p
    
    def _add_normal_paragraph(self, text: str = "", indent: bool = False):
        """Add normal paragraph"""
        p = self.doc.add_paragraph(text, style='Normal')
        if indent:
            p.paragraph_format.left_indent = Inches(0.5)
        return p
    
    def _set_cell_background(self, cell, color: str):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def _add_critical_warning_box(self, title: str, content: List[str]):
        """Add a warning/information box"""
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        cell = table.rows[0].cells[0]
        self._set_cell_background(cell, 'FFE699')  # Light yellow
        
        p = cell.paragraphs[0]
        run = p.add_run(f"⚠ {title}: ")
        run.bold = True
        
        for item in content:
            p.add_run(f"\n• {item}")
        
        return table
    
    def generate(self) -> Document:
        """Generate the complete DPA document"""
        self._generate_title_page()
        self.doc.add_page_break()
        self._generate_preamble()
        self.doc.add_page_break()
        self._generate_articles()
        self.doc.add_page_break()
        self._generate_appendices()
        
        return self.doc
    
    def _generate_title_page(self):
        """Generate title page"""
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("GDPR DATA PROCESSOR AGREEMENT")
        run.bold = True
        run.font.size = Pt(16)
        
        # Service name
        service = self.doc.add_paragraph()
        service.alignment = WD_ALIGN_PARAGRAPH.CENTER
        service_run = service.add_run(f"For: {self.data.service_name}")
        service_run.font.size = Pt(12)
        
        self.doc.add_paragraph()  # Spacing
        
        # Parties section
        self._add_heading("Parties", level=2)
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph("Data Controller", style='Heading 3')
        self._add_normal_paragraph(f"Name: {self.data.controller_name}", indent=True)
        self._add_normal_paragraph(f"CVR/Tax ID: {self.data.controller_cvr}", indent=True)
        self._add_normal_paragraph(f"Address: {self.data.controller_address}", indent=True)
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph("Data Processor", style='Heading 3')
        self._add_normal_paragraph(f"Name: {self.data.processor_name}", indent=True)
        self._add_normal_paragraph(f"CVR/Tax ID: {self.data.processor_cvr}", indent=True)
        self._add_normal_paragraph(f"Address: {self.data.processor_address}", indent=True)
        
        self.doc.add_paragraph()
        
        # Key information
        self._add_heading("Key Information", level=2)
        
        info_table = self.doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        rows = info_table.rows
        cells = rows[0].cells
        cells[0].text = "Execution Date"
        cells[1].text = self.data.execution_date
        
        cells = rows[1].cells
        cells[0].text = "Storage Location"
        cells[1].text = self.data.storage_location
        
        cells = rows[2].cells
        cells[0].text = "Governing Jurisdiction"
        cells[1].text = self.data.country
        
        cells = rows[3].cells
        cells[0].text = "Service Description"
        cells[1].text = self.data.service_name
        
        self.doc.add_paragraph()
        
        # Health data warning
        if self.health_data_detected:
            self._add_critical_warning_box(
                "HEALTH DATA DETECTED",
                [
                    f"Reason: {self.health_data_reason}",
                    "Enhanced security requirements apply (Article 9 GDPR)",
                    "Explicit consent may be required",
                    "Data Protection Impact Assessment (DPIA) is mandatory",
                    "Annual security audits are recommended"
                ]
            )
        
        self.doc.add_paragraph()
        self._add_normal_paragraph(
            "This Data Processor Agreement (DPA) is executed as of the date above and governs the processing "
            "of personal data according to Article 28(3) of the General Data Protection Regulation (GDPR)."
        )
    
    def _generate_preamble(self):
        """Generate preamble section"""
        self._add_heading("Preamble", level=1)
        
        # Reference to GDPR
        country_mod = CountryModifier()
        preamble_text = country_mod.get_modification(self.data.country, "preamble_addition")
        
        self._add_normal_paragraph(
            f"For the purposes of Article 28(3) of Regulation (EU) 2016/679 (the GDPR) {preamble_text},"
        )
        
        self._add_normal_paragraph()
        
        self._add_normal_paragraph(
            f"This Data Processor Agreement (the 'Agreement') is entered into between "
            f"{self.data.controller_name} (the 'Data Controller') and "
            f"{self.data.processor_name} (the 'Data Processor')."
        )
        
        self._add_normal_paragraph()
        
        self._add_normal_paragraph(
            "Whereas:\n\n"
            "(A) The Data Controller and Data Processor wish to ensure that the processing of personal data "
            "is carried out in accordance with the GDPR and applicable national data protection legislation.\n\n"
            "(B) The Data Processor shall provide sufficient guarantees as to the implementation of appropriate "
            "technical and organizational measures.\n\n"
            "(C) The purpose of this Agreement is to set out the respective rights and obligations of the parties "
            "and to comply with the requirements of Article 28(3) GDPR."
        )
        
        self._add_normal_paragraph()
        
        self._add_normal_paragraph(
            "The parties have agreed to enter into this Data Processor Agreement on the terms and conditions set out below."
        )
        
        # Authority reference
        authority = country_mod.get_modification(self.data.country, "authority")
        self._add_normal_paragraph()
        self._add_normal_paragraph(
            f"This Agreement is subject to oversight by {authority}."
        )
    
    def _generate_articles(self):
        """Generate the core article clauses"""
        self._add_heading("Terms and Conditions", level=1)
        
        articles = [
            {
                "number": 1,
                "title": "Definitions",
                "content": [
                    "'Personal Data' means any information relating to an identified or identifiable natural person ('Data Subject');",
                    "'Processing' means any operation performed on Personal Data, such as collection, recording, organization, structuring, storage, adaptation, retrieval, consultation, use, disclosure, transmission, or erasure;",
                    "'Special Category Data' means Personal Data revealing racial or ethnic origin, political opinions, religious or philosophical beliefs, trade union membership, genetic data, biometric data for identification purposes, health data, or data concerning sex life or sexual orientation;",
                    "'Data Breach' means a breach of security leading to the accidental or unlawful destruction, loss, alteration, unauthorized disclosure of, or access to Personal Data;",
                    "'Sub-processor' means any natural or legal person processing Personal Data on behalf of the Data Processor."
                ]
            },
            {
                "number": 2,
                "title": "Subject Matter and Scope of Processing",
                "content": [
                    f"Service Description: {self.data.service_name}",
                    f"Data Categories: {self.data.data_categories or (', '.join(self.data.data_types) if self.data.data_types else 'As specified in Appendix A')}",
                    f"Processing Purposes: {self.data.purposes or 'Service delivery and related activities'}",
                    f"Storage Location: {self.data.storage_location}"
                ]
            },
            {
                "number": 3,
                "title": "Processor Obligations",
                "content": [
                    "Process Personal Data only on documented instructions from the Data Controller;",
                    "Ensure that persons authorized to process Personal Data are bound by confidentiality or under an appropriate legal obligation of confidentiality;",
                    "Implement appropriate technical and organizational measures to ensure security of Personal Data;",
                    "Not disclose or permit access to Personal Data except on instruction of the Data Controller or as required by law;",
                    "Assist the Data Controller in fulfilling Data Subject rights requests;",
                    "Not engage Sub-processors without prior specific or general written authorization from the Data Controller;",
                    "Impose data protection obligations on Sub-processors equivalent to those in this Agreement;",
                    "Assist the Data Controller in ensuring compliance with GDPR Articles 32-36 (security, DPIA, breach notification)."
                ]
            },
            {
                "number": 4,
                "title": "Sub-processors",
                "content": [
                    f"The Data Processor shall notify the Data Controller of any intended changes regarding the engagement or replacement of Sub-processors, with notice of at least {CountryModifier.get_modification(self.data.country, 'sub_processor_notice')}.;",
                    "The Data Controller may object to the engagement of a new Sub-processor on reasonable grounds relating to data protection compliance.;",
                    "If the Data Controller objects, the parties shall negotiate in good faith with the aim of resolving the objection within 10 business days.;",
                    "Should a mutually acceptable solution not be found, the Data Controller may terminate the Agreement without penalty.",
                    "The Data Processor shall ensure Sub-processors are contractually bound by equivalent data protection obligations."
                ]
            },
            {
                "number": 5,
                "title": "Data Security",
                "content": [
                    "The Data Processor shall implement and maintain technical and organizational measures appropriate to the risk level of the Personal Data being processed;",
                    "Security measures shall include: encryption (AES-256 minimum for data at rest, TLS 1.2+ for data in transit), access controls, authentication, audit logging, and intrusion detection;",
                    "For Special Category Data or Health Data, enhanced security measures shall be implemented including multi-factor authentication, quarterly access reviews, and annual security audits;",
                    "The Data Processor shall conduct annual risk assessments and security audits (SOC 2 Type II or ISO 27001 certification recommended)."
                ]
            },
            {
                "number": 6,
                "title": "Data Breach Notification",
                "content": [
                    f"The Data Processor shall notify the Data Controller {CountryModifier.get_modification(self.data.country, 'breach_notification')} of becoming aware of a Personal Data Breach;",
                    "The notification shall include: (a) description of the breach; (b) likely consequences; (c) data categories and approximate number of Data Subjects affected; (d) measures taken or proposed to remedy the breach;",
                    "The Data Processor shall provide reasonable assistance to the Data Controller in meeting breach notification obligations to Data Subjects and supervisory authorities;",
                    "The Data Processor shall preserve all evidence related to the breach for investigation and potential legal proceedings."
                ]
            },
            {
                "number": 7,
                "title": "Data Subject Rights",
                "content": [
                    "The Data Processor shall assist the Data Controller in fulfilling Data Subject requests for access, rectification, erasure, restriction, portability, and objection under GDPR Articles 12-21;",
                    "The Data Processor shall respond to Data Subject requests within 15 business days of receiving them from the Data Controller;",
                    "The Data Processor shall not directly respond to Data Subjects except as instructed by the Data Controller;",
                    "The Data Processor shall provide all necessary information and cooperation to enable the Data Controller to demonstrate compliance with GDPR obligations."
                ]
            },
            {
                "number": 8,
                "title": "Retention and Deletion",
                "content": [
                    f"Personal Data shall be processed and retained for the duration of this Agreement and for a period of {self.data.retention_period} thereafter;",
                    "Upon termination of this Agreement, the Data Processor shall, at the Data Controller's election: (a) return all Personal Data; or (b) delete all Personal Data (including backups) within 30 days;",
                    "The Data Processor shall demonstrate deletion or return with written certification;",
                    "Data retention for legal or compliance purposes may continue only where required by law and shall be properly documented.",
                    "Deleted data shall be permanently erased and not recoverable from backups."
                ]
            },
            {
                "number": 9,
                "title": "Audit and Compliance",
                "content": [
                    "The Data Processor shall make available all information necessary to demonstrate compliance with this Agreement and GDPR;",
                    "The Data Controller may conduct audits and inspections of the Data Processor's processing activities;",
                    "The Data Processor shall provide audit reports (SOC 2, ISO 27001, or equivalent) to the Data Controller upon request;",
                    "Audits may be conducted annually or more frequently if required by the Data Controller;",
                    "The Data Processor shall remediate any identified non-compliance within agreed timeframes."
                ]
            },
            {
                "number": 10,
                "title": "Data Transfer Outside EEA",
                "content": [
                    "Personal Data shall not be transferred outside the EU/EEA unless the Data Controller provides explicit prior written approval;",
                    "Transfers to adequate countries may proceed under derogations specified by GDPR Article 49;",
                    "Transfers to non-adequate countries require Standard Contractual Clauses (SPCs) or equivalent mechanism;",
                    "The Data Processor shall notify the Data Controller immediately of any government requests for Personal Data."
                ]
            },
            {
                "number": 11,
                "title": "Liability",
                "content": [
                    "The Data Processor is liable to the Data Controller for any damages caused by breach of GDPR obligations under this Agreement;",
                    "The Data Processor shall indemnify and hold harmless the Data Controller against third-party claims arising from the Data Processor's breach of this Agreement;",
                    "The Data Processor's total liability is not limited with respect to: (a) unauthorized disclosure of Personal Data; (b) failure to implement mandatory security measures; or (c) violation of GDPR Article 9 (Special Category Data);",
                    "For other breaches, liability may be capped at annual fees paid under this Agreement, provided the Data Processor has appropriate insurance."
                ]
            },
            {
                "number": 12,
                "title": "Term and Termination",
                "content": [
                    "This Agreement shall commence on the date specified above and shall continue until terminated by either party;",
                    "Either party may terminate for material breach if such breach is not cured within 15 calendar days of written notice;",
                    "The Data Controller may terminate immediately if the Data Processor fails to comply with GDPR or this Agreement in a way that significantly impacts Data Subject rights;",
                    "Upon termination, the Data Processor shall comply with the deletion/return obligations in Article 8."
                ]
            },
            {
                "number": 13,
                "title": "Governing Law and Jurisdiction",
                "content": [
                    f"This Agreement shall be governed by the laws of {self.data.country};",
                    f"The parties agree to the jurisdiction of the courts in {self.data.country} for any disputes;",
                    "The parties agree to attempt resolution of disputes through good faith negotiation before initiating legal proceedings."
                ]
            }
        ]
        
        for article in articles:
            # Article heading
            p = self.doc.add_paragraph(f"Article {article['number']}: {article['title']}", style='Heading 2')
            
            # Article content
            for item in article['content']:
                self.doc.add_paragraph(item, style='List Paragraph')
            
            self.doc.add_paragraph()  # Spacing
    
    def _generate_appendices(self):
        """Generate appendices"""
        self._add_heading("Appendices", level=1)
        
        # Appendix A: Details of Processing
        self._add_heading("Appendix A: Details of Processing", level=2)
        
        appendix_a_table = self.doc.add_table(rows=8, cols=2)
        appendix_a_table.style = 'Light Grid Accent 1'
        
        rows = appendix_a_table.rows
        rows[0].cells[0].text = "Data Categories"
        rows[0].cells[1].text = ", ".join(self.data.data_types) if self.data.data_types else "As specified by Data Controller"
        
        rows[1].cells[0].text = "Special Categories"
        rows[1].cells[1].text = "YES - Health Data Detected" if self.health_data_detected else "None identified"
        
        rows[2].cells[0].text = "Processing Purposes"
        rows[2].cells[1].text = self.data.purposes or "Service delivery"
        
        rows[3].cells[0].text = "Data Subjects"
        rows[3].cells[1].text = "As specified in instructions"
        
        rows[4].cells[0].text = "Retention Period"
        rows[4].cells[1].text = self.data.retention_period
        
        rows[5].cells[0].text = "Primary Storage Location"
        rows[5].cells[1].text = self.data.storage_location
        
        rows[6].cells[0].text = "Data Recipient Categories"
        rows[6].cells[1].text = "Data Processor staff, authorized auditors, Sub-processors (if approved)"
        
        rows[7].cells[0].text = "Processing Location(s)"
        rows[7].cells[1].text = self.data.storage_location
        
        self.doc.add_paragraph()
        
        # Appendix B: Technical and Organizational Measures
        self._add_heading("Appendix B: Technical and Organizational Measures (TOMs)", level=2)
        
        toms_content = [
            ("Encryption", [
                "Data at rest: AES-256 or equivalent encryption",
                "Data in transit: TLS 1.2 minimum (TLS 1.3 recommended)",
                "Encryption keys: Managed separately from encrypted data"
            ]),
            ("Access Controls", [
                "Role-based access control (RBAC)",
                "Principle of least privilege",
                "Multi-factor authentication for privileged access",
                "Quarterly access reviews and removal of unused access"
            ]),
            ("Audit and Monitoring", [
                "Immutable audit logs for all data access",
                "Log retention: Minimum 3 years for health data, 1 year otherwise",
                "Real-time alerting for unauthorized access attempts",
                "Monthly audit log review"
            ]),
            ("Data Isolation", [
                "Separate database/schema for sensitive data",
                "Network segmentation where feasible",
                "Isolated backup storage"
            ]),
            ("Backup and Recovery", [
                "Daily backups with encrypted storage",
                "Recovery time objective (RTO): Maximum 24 hours",
                "Recovery point objective (RPO): Maximum 1 day",
                "Regular restoration testing (quarterly)"
            ]),
            ("Personnel Security", [
                "Data protection training for all staff (annual minimum)",
                "Background checks for staff accessing sensitive data",
                "Confidentiality agreements upon hire",
                "Exit procedures and access revocation"
            ]),
            ("Vulnerability Management", [
                "Quarterly vulnerability scanning",
                "Annual penetration testing",
                "Timely patching (critical patches within 48 hours)",
                "Security incident response plan"
            ]),
            ("Compliance Audits", [
                "Third-party security audit (SOC 2 Type II, ISO 27001, or equivalent)",
                "Audit frequency: Annual (or as required)",
                "Results provided to Data Controller within 30 days"
            ])
        ]
        
        for measure_name, details in toms_content:
            p = self.doc.add_paragraph(f"{measure_name}:", style='Heading 3')
            for detail in details:
                self.doc.add_paragraph(detail, style='List Paragraph')
            self.doc.add_paragraph()
        
        # Appendix C: Sub-processors and Recipients
        self._add_heading("Appendix C: Sub-processors and Third-Party Recipients", level=2)
        
        self._add_normal_paragraph(
            "The following Sub-processors are authorized to process Personal Data under this Agreement:"
        )
        
        if self.data.sub_processors:
            sub_table = self.doc.add_table(rows=len(self.data.sub_processors) + 1, cols=4)
            sub_table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = sub_table.rows[0].cells
            header_cells[0].text = "Sub-processor Name"
            header_cells[1].text = "Location"
            header_cells[2].text = "Processing Activities"
            header_cells[3].text = "Approval Date"
            
            for idx, sp in enumerate(self.data.sub_processors, 1):
                cells = sub_table.rows[idx].cells
                cells[0].text = sp.get("name", "TBD")
                cells[1].text = sp.get("location", "TBD")
                cells[2].text = sp.get("activities", "TBD")
                cells[3].text = sp.get("approval_date", datetime.now().strftime("%Y-%m-%d"))
        else:
            self._add_normal_paragraph(
                "Sub-processors to be designated by separate notice. Any new Sub-processors will require "
                "Data Controller approval as specified in Article 4."
            )
        
        self.doc.add_paragraph()
        
        # Appendix D: Data Controller Instructions
        self._add_heading("Appendix D: Data Controller Instructions", level=2)
        
        self._add_normal_paragraph(
            "The Data Processor shall process Personal Data only in accordance with the following instructions from the Data Controller:"
        )
        
        instructions = [
            f"Process only {self.data.data_categories or 'specified data categories'} for the purposes of {self.data.purposes or 'service delivery'};",
            f"Store data primarily in {self.data.storage_location} with no transfers outside the EEA without explicit approval;",
            "Implement all security measures specified in Appendix B;",
            "Respond to Data Subject rights requests within 15 business days;",
            "Notify the Data Controller of Data Breaches within the timeframes specified in Article 6;",
            "Provide audit access and reports as requested by the Data Controller;",
            "Notify of any regulatory investigations or legal requests affecting this processing;",
            "Delete all Personal Data within 30 days of contract termination or upon Data Controller request;",
            "Cooperate fully with Data Protection Impact Assessments (DPIAs) if required;",
            f"Comply with retention periods specified in Appendix A ({self.data.retention_period})."
        ]
        
        for instruction in instructions:
            self.doc.add_paragraph(instruction, style='List Paragraph')
        
        self.doc.add_paragraph()
        
        # Appendix E: Contact Information and Notifications
        self._add_heading("Appendix E: Contact Information and Notifications", level=2)
        
        contact_table = self.doc.add_table(rows=4, cols=2)
        contact_table.style = 'Light Grid Accent 1'
        
        contact_rows = contact_table.rows
        contact_rows[0].cells[0].text = "Data Controller Contact"
        contact_rows[0].cells[1].text = self.data.controller_name
        
        contact_rows[1].cells[0].text = "Data Processor Contact"
        contact_rows[1].cells[1].text = self.data.processor_name
        
        contact_rows[2].cells[0].text = "Breach Notification Address"
        contact_rows[2].cells[1].text = "[To be provided by Data Controller]"
        
        contact_rows[3].cells[0].text = "Legal Inquiries"
        contact_rows[3].cells[1].text = "[To be provided by both parties]"
        
        self.doc.add_page_break()
        
        # Signature page
        self._add_heading("Execution", level=2)
        
        self._add_normal_paragraph(
            f"This Data Processor Agreement is executed as of {self.data.execution_date} and shall be effective as of this date."
        )
        
        self.doc.add_paragraph("\n\n")
        
        # Signature blocks
        self._add_normal_paragraph("DATA CONTROLLER:")
        self.doc.add_paragraph("_" * 50)
        self._add_normal_paragraph(f"{self.data.controller_name}")
        self._add_normal_paragraph("Authorized Representative")
        
        self.doc.add_paragraph("\n\n")
        
        self._add_normal_paragraph("DATA PROCESSOR:")
        self.doc.add_paragraph("_" * 50)
        self._add_normal_paragraph(f"{self.data.processor_name}")
        self._add_normal_paragraph("Authorized Representative")
        
        self.doc.add_paragraph("\n")
        self._add_normal_paragraph(f"Date: {self.data.execution_date}")


def generate_dpa(data: AgreementData, output_path: str = "DPA_Agreement.docx") -> str:
    """
    Generate a GDPR Data Processor Agreement
    
    Args:
        data: AgreementData object with all required information
        output_path: Path where to save the generated document
    
    Returns:
        Path to the generated document
    """
    generator = DPAGenerator(data)
    doc = generator.generate()
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    # Example usage
    example_data = AgreementData(
        controller_name="Example Corp A/S",
        controller_cvr="12345678",
        controller_address="Commercial Street 1, Copenhagen, Denmark",
        processor_name="Data Service Solutions Ltd",
        processor_cvr="87654321",
        processor_address="Tech Avenue 42, Dublin, Ireland",
        service_name="Cloud Data Processing Platform",
        data_types=["Names", "Email addresses", "Phone numbers", "Location data"],
        storage_location="Ireland (EU)",
        country="Germany",
        purposes="Service delivery and analytics",
        data_categories="Customer contact information and platform usage data",
        execution_date=datetime.now().strftime("%B %d, %Y")
    )
    
    output = generate_dpa(example_data, "example_dpa.docx")
    print(f"Generated DPA at: {output}")
